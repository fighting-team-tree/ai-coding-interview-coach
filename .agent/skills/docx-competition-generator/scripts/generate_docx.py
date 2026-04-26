import docx
import re
import argparse
import os

def clean_markdown(text):
    if not text:
        return ""
    # 1. 볼드/이탤릭 제거
    text = re.sub(r'\*\*+(.*?)\*\*+', r'\1', text)
    text = re.sub(r'\*+(.*?)\*+', r'\1', text)
    text = re.sub(r'__+(.*?)__+', r'\1', text)
    
    # 2. Markdown 헤더 제거
    text = re.sub(r'#+\s*(.*?)\n', r'\1\n', text)
    
    # 3. <br> 태그를 실제 줄바꿈으로 변환
    text = text.replace("<br>", "\n")
    
    # 4. 특수 체크박스 기호 보정 ( -> ■)
    text = text.replace("", "■")
    
    # 5. 불필요한 공백 정리
    text = text.strip()
    return text

def get_md_value(tag, md_content):
    # 정규표현식으로 표 안의 값 추출 (멀티라인 고려)
    pattern = fr"\| {re.escape(tag)} \| (.*?) \|"
    match = re.search(pattern, md_content, re.DOTALL)
    if match:
        val = match.group(1)
        return clean_markdown(val)
    return ""

def generate_docx(md_path, template_path, output_path):
    if not os.path.exists(md_path):
        print(f"Error: MD file not found at {md_path}")
        return
    if not os.path.exists(template_path):
        print(f"Error: Template not found at {template_path}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = docx.Document(template_path)

    # Table 1: 프로젝트 정보 (상단 요약 - 지원트랙 및 인원)
    if len(doc.tables) > 1:
        summary_info = doc.tables[1]
        try:
            for row in summary_info.rows:
                if len(row.cells) < 2: continue
                c0_text = row.cells[0].text.strip()
                if "지원트랙" in c0_text:
                    md_val = get_md_value("지원트랙", md_content)
                    row.cells[1].text = md_val if md_val else "[ √ ] 국내 AI 트랙"
                elif "팀인원" in c0_text:
                    md_val = get_md_value("팀인원 수", md_content)
                    row.cells[1].text = md_val if md_val else "4명"
        except Exception as e:
            print(f"Warning: Failed to fill summary table: {e}")

    # 날짜 치환: Markdown에서 YYYY. MM. DD. 형식의 날짜를 찾아 Word의 '2026. xx. xx.' 부분을 교체
    date_match = re.search(r'(20\d\d\.\s*\d{1,2}\.\s*\d{1,2}\.)', md_content)
    if date_match:
        submit_date = date_match.group(1)
        for p in doc.paragraphs:
            if "xx. xx." in p.text:
                p.text = p.text.replace("2026. xx. xx.", submit_date).replace("202x. xx. xx.", submit_date)

    # Markdown에서 모든 필드 미리 추출
    all_fields = {}
    rows = re.findall(r"\| (.*?) \| (.*?) \|", md_content)
    for k, v in rows:
        all_fields[k.strip()] = clean_markdown(v)

    # 매핑 데이터 정의 (공식 템플릿의 테이블 인덱스 기반)
    sections = [
        {"idx": 3, "prefix": "1."},
        {"idx": 5, "prefix": "2."},
        {"idx": 7, "prefix": "3."},
        {"idx": 9, "prefix": "4."}
    ]

    for section in sections:
        if len(doc.tables) > section["idx"]:
            table = doc.tables[section["idx"]]
            for row in table.rows:
                if len(row.cells) < 2: continue
                key_text = row.cells[0].text.strip()
                # 1.1, 1.2, 1.2.1 등 모든 패턴 추출
                match = re.search(fr"({re.escape(section['prefix'])}\d+(\.\d+)*)", key_text)
                if match:
                    tag_id = match.group(1)
                    found = False
                    # 1. Exact match first
                    for field_key, field_val in all_fields.items():
                        if field_key.startswith(tag_id):
                            # 하드웨어 예/아니오
                            row.cells[1].text = field_val
                            
                            # 삽입된 텍스트의 단락 공백(줄간격, 문단 앞뒤 간격)을 강제로 최소화하여 페이지를 줄임
                            from docx.shared import Pt
                            for paragraph in row.cells[1].paragraphs:
                                if paragraph.paragraph_format is not None:
                                    paragraph.paragraph_format.space_after = Pt(0)
                                    paragraph.paragraph_format.space_before = Pt(0)
                                    paragraph.paragraph_format.line_spacing = 1.0
                            
                            found = True
                            break
                    if not found and "하드웨어 설명" in key_text:
                        row.cells[1].text = "해당 없음"
                elif "지원트랙" in key_text:
                    row.cells[1].text = "[ √ ] 국내 AI 트랙"

    # 참고 테이블 및 안내 테이블 삭제
    tables_to_remove = []
    for i, t in enumerate(doc.tables):
        try:
            first_cell = t.cell(0,0).text.strip() if len(t.rows) > 0 and len(t.columns) > 0 else ""
            # 참고 글상자 또는 작성방법 설명 삭제
            if "<참고" in first_cell or "작성방법" in first_cell or "본 문서는 AI 챔피언" in first_cell:
                tables_to_remove.append(t)
            # 내용 없이 빈 셀만 있는 '참고'용 구조 삭제
            if i % 2 == 0 and i > 0 and not first_cell:
                 tables_to_remove.append(t)
        except:
            continue
    
    for t in reversed(tables_to_remove):
        try:
            p = t._element.getparent()
            p.remove(t._element)
        except:
            pass

    # 섹션 1 (1. 일반 사항) 시작 전에 페이지 브레이크 삽입 (표지와 분리)
    for t in doc.tables:
        if len(t.rows) > 0 and len(t.columns) > 0:
            if "1.1 기술명" in t.cell(0,0).text:
                # 이미 페이지 브레이크가 있는지 확인
                prev = t._element.getprevious()
                if prev is not None and 'w:br' in prev.xml and 'w:type="page"' in prev.xml:
                    continue
                new_para = doc.add_paragraph()
                new_para.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
                t._element.addprevious(new_para._element)
                break

    # 불필요한 단락 제거 (안내 문구 및 가이드 박스 내용)
    paras_to_remove = []
    summary_table_passed = False
    
    keywords = [
        "본 문서는 AI 챔피언", 
        "작성방법에 대한 설명은", 
        "제출 시 양식의 작성방법", 
        "구현제안서는 표지 제외 5장 이내",
        "초과하는 부분은 평가에서 제외됨"
    ]

    for child in doc.element.body:
        if child.tag.endswith('tbl'):
            from docx.table import Table
            t = Table(child, doc)
            if len(t.rows) > 0 and "지원트랙" in t.cell(0,0).text:
                summary_table_passed = True
        
        elif child.tag.endswith('p'):
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc)
            text = para.text.strip()
            
            # 1. 안내 문구 강제 삭제 (표지/본문 불문)
            if any(k in text for k in keywords):
                paras_to_remove.append(para)
                continue
                
            # 2. 표지 요약 테이블 이후의 빈 줄 삭제 (단, 페이지 브레이크 보호)
            if summary_table_passed and not text:
                if 'w:br' in para._element.xml and 'w:type="page"' in para._element.xml:
                    continue
                paras_to_remove.append(para)
    
    for p in paras_to_remove:
        try:
            p._element.getparent().remove(p._element)
        except:
            pass
    
    for p in paras_to_remove:
        try:
            p._element.getparent().remove(p._element)
        except:
            pass

    # 최종 저장
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate competition DOCX from Markdown.")
    parser.add_argument("--input", required=True, help="Path to input Markdown file")
    parser.add_argument("--template", required=True, help="Path to template DOCX file")
    parser.add_argument("--output", required=True, help="Path to save output DOCX")
    
    args = parser.parse_args()
    generate_docx(args.input, args.template, args.output)
